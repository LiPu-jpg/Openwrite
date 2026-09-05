from __future__ import annotations

import json
import os
import hashlib
import subprocess
import tempfile
import unittest
from pathlib import Path

from dog_import import write_import_artifacts
from dog_delivery import write_delivery_artifacts
from dog_review import build_review_manifest, write_review_artifacts
from smart_import import detect_and_convert
from tools.review_dag_framework import instantiate_review_dag, review_dag_framework


REPO_ROOT = Path(__file__).resolve().parent.parent
VERIFIERS = REPO_ROOT / "scripts" / "dog"


class _Studio:
    def __init__(self, root: Path, novel_id: str) -> None:
        self.root = root
        self.novel_id = novel_id

    def get(self, path: str) -> dict:
        if path != "/api/workspace":
            raise AssertionError(path)
        return {
            "project": {"root": str(self.root)},
            "snapshot": {"novel_id": self.novel_id},
        }


def _run_verifier(name: str, target: Path) -> dict:
    result = subprocess.run(
        ["node", str(VERIFIERS / f"{name}.js"), str(target)],
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def _validate_with_dog(graph_path: Path, workspace: Path) -> None:
    dog_dir = os.environ.get("DSH_DOG_DIR", "").strip()
    if not dog_dir:
        return
    dog_root = Path(dog_dir).resolve()
    core = dog_root / "lib" / "core.js"
    if not core.is_file():
        raise AssertionError(f"dsh-dog build missing: {core}")
    source = """
import fs from 'node:fs'
import { execFile } from 'node:child_process'
import { promisify } from 'node:util'
import { pathToFileURL } from 'node:url'
import { join } from 'node:path'
const dogRoot = process.argv[1]
const chunk = fs.readdirSync(join(dogRoot, 'lib')).find(name => /^core-.*\\.js$/.test(name))
if (!chunk) throw new Error('dsh-dog internal core bundle is missing')
const internal = await import(pathToFileURL(join(dogRoot, 'lib', chunk)).href)
const DogRepository = internal.a
const DogEngine = internal.t
const resolveScriptPath = internal.n
if (!DogRepository || !DogEngine || !resolveScriptPath) throw new Error('dsh-dog core exports changed')
const graph = JSON.parse(fs.readFileSync(process.argv[2], 'utf8'))
const repository = new DogRepository(join(process.argv[3], `.dog-test-${Date.now()}`))
const runFile = promisify(execFile)
const engine = new DogEngine({
  config: {
    workspaceRoot: process.argv[3], scriptsDirectory: process.argv[4], storageDirectory: 'dog',
    maxGraphNodes: 256, maxExpressionNodes: 512, maxExpressionDepth: 32,
    maxSandboxBytes: 10485760, allowPartialRoot: false,
    maxConcurrentVerifications: 4, revalidateThreshold: 0.5, gmDigestAlgo: 'sha256',
  },
  repository,
  programmatic: async (script, inputPath) => {
    const result = await runFile(process.execPath, [resolveScriptPath(process.argv[4], script), inputPath])
    const value = JSON.parse(result.stdout)
    return { state: value.verdict, evidence: value.evidence }
  },
  agentic: async () => ({ state: 'pass', evidence: { contractSmoke: true } }),
})
const report = engine.validate(graph)
if (!report.valid) throw new Error(JSON.stringify(report))
await engine.create(graph)
const run = await engine.run(graph.id)
if (!['completed', 'success', 'failure', 'needs_replan', 'partial_success'].includes(run.state)) {
  throw new Error(`DoG run did not settle: ${run.state}`)
}
"""
    result = subprocess.run(
        [
            "node", "--input-type=module", "-e", source, str(dog_root),
            str(graph_path), str(workspace), str(VERIFIERS),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        raise AssertionError(f"dsh-dog execution failed:\n{result.stderr or result.stdout}")


def _assert_acyclic(graph: dict) -> None:
    adjacency: dict[str, list[str]] = {node_id: [] for node_id in graph["nodes"]}
    for edge in graph.get("dependsOn", []):
        adjacency[edge["target"]].append(edge["source"])
    visiting: set[str] = set()
    visited: set[str] = set()

    def visit(node_id: str) -> None:
        if node_id in visiting:
            raise AssertionError(f"cycle found at {node_id}")
        if node_id in visited:
            return
        visiting.add(node_id)
        for child in adjacency[node_id]:
            visit(child)
        visiting.remove(node_id)
        visited.add(node_id)

    for node_id in adjacency:
        visit(node_id)


class DogArtifactTests(unittest.TestCase):
    def test_review_graph_preserves_partial_and_hard_verdicts(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            review = {
                "passed": False,
                "score": 82,
                "summary": "存在一项硬错误",
                "dimensions": [1, 2],
                "issue_details": [
                    {
                        "id": "issue-1", "dimension": 1, "severity": "critical",
                        "description": "角色行为违背设定", "evidence": {"quote": "证据"},
                    },
                    {
                        "id": "issue-general", "dimension": None, "severity": "warning",
                        "description": "无法归入维度的问题", "evidence": {"quote": "旁证"},
                    },
                ],
            }
            result = write_review_artifacts(_Studio(root, "book"), "ch_001", review, 70)
            graph_path = Path(result["graph_path"])
            artifact_dir = graph_path.parent

            self.assertEqual(len(result["graph"]["nodes"]), 47)
            self.assertFalse(any(
                node.get("verifier", {}).get("mode") == "agentic"
                for node in result["graph"]["nodes"].values()
            ))
            parents = {
                edge["child"]: edge["parent"]
                for edge in result["graph"]["contains"]
            }
            self.assertEqual(parents["domain-character"], "root")
            self.assertEqual(parents["dim-01"], "domain-character")
            self.assertEqual(parents["dim-27"], "gate")
            self.assertEqual(_run_verifier("review-dimension", artifact_dir / "dim_01.json")["verdict"], "fail")
            self.assertEqual(_run_verifier("review-dimension", artifact_dir / "dim_02.json")["verdict"], "pass")
            self.assertEqual(_run_verifier("review-dimension", artifact_dir / "dim_03.json")["verdict"], "inconclusive")
            manifest = json.loads((artifact_dir / "review.json").read_text(encoding="utf-8"))
            self.assertEqual(manifest["issueCount"], 2)
            self.assertEqual(manifest["unmappedIssueCount"], 1)
            self.assertEqual(manifest["schemaVersion"], "dsh-novel.review.manifest.v2")
            self.assertEqual(manifest["frameworkId"], "openwrite.standard-chapter-review")
            self.assertEqual(manifest["frameworkRevision"], result["framework"]["revision"])
            self.assertEqual(len(manifest["domains"]), 6)
            self.assertEqual(_run_verifier("review-record", artifact_dir / "gate.json")["verdict"], "fail")
            _assert_acyclic(result["graph"])
            _validate_with_dog(graph_path, root)

    def test_python_and_typescript_instantiate_the_same_canonical_framework(self) -> None:
        framework = review_dag_framework()
        relative_dir = "data/novels/book/data/dog/reviews/ch_007"
        python_graph = instantiate_review_dag("ch_007", relative_dir, framework=framework)
        script = """
import fs from 'node:fs'
import { instantiateDogReviewFramework } from './packages/openwrite-bridge/lib/dog-review.js'
const framework = JSON.parse(fs.readFileSync(0, 'utf8'))
process.stdout.write(JSON.stringify(instantiateDogReviewFramework(
  framework, 'ch_007', 'data/novels/book/data/dog/reviews/ch_007'
)))
"""
        result = subprocess.run(
            ["node", "--input-type=module", "-e", script],
            cwd=REPO_ROOT,
            input=json.dumps(framework, ensure_ascii=False),
            capture_output=True,
            text=True,
            check=True,
        )
        self.assertEqual(json.loads(result.stdout), python_graph)

    def test_python_and_typescript_review_contracts_match(self) -> None:
        review = {
            "score": 84,
            "passed": False,
            "source_revision": "sha256:test",
            "issue_details": [{
                "id": "issue-1", "dimension": 1, "severity": "critical",
                "revision_priority": "blocker", "description": "人物失真", "evidence": "证据",
            }],
            "review_v2": {
                "schema_version": "openwrite.review.v2",
                "execution_status": "partial", "quality_score": 84, "coverage": 0.9,
                "gate_status": "blocked", "delivery_status": "blocked",
                "production_gate_status": "disabled_uncalibrated",
                "requested_dimensions": list(range(1, 38)),
                "domains": [{
                    "id": "character", "status": "evaluated", "earned": 12, "max": 15,
                    "potential_max": 15, "coverage": 1,
                    "criteria": [{
                        "id": "character_fidelity", "status": "evaluated", "earned": 4,
                        "max": 5, "evidence": ["证据"], "legacy_check_ids": [1, 34],
                    }],
                }],
            },
        }
        python_manifest, python_dimensions = build_review_manifest(review, "ch_001", 70)
        script = """
import fs from 'node:fs'
import { buildDogReviewBundle } from './packages/openwrite-bridge/lib/dog-review.js'
const review = JSON.parse(fs.readFileSync(0, 'utf8'))
process.stdout.write(JSON.stringify(buildDogReviewBundle(review, 'ch_001', 70)))
"""
        result = subprocess.run(
            ["node", "--input-type=module", "-e", script],
            cwd=REPO_ROOT,
            input=json.dumps(review),
            capture_output=True,
            text=True,
            check=True,
        )
        typescript = json.loads(result.stdout)
        typescript_manifest = typescript["manifest"]
        for key in (
            "schemaVersion", "chapterId", "verdict", "executionStatus", "qualityScore",
            "coverage", "gateStatus", "deliveryStatus", "requestedDimensions", "dimensionCount",
            "issueCount", "unmappedIssueCount",
        ):
            self.assertEqual(typescript_manifest[key], python_manifest[key], key)
        self.assertEqual(
            [(item["id"], item["legacyCheckIds"], item["verdict"]) for item in typescript_manifest["domains"]],
            [(item["id"], item["legacyCheckIds"], item["verdict"]) for item in python_manifest["domains"]],
        )
        self.assertEqual(
            [(item["dimension"], item["status"], item["verdict"]) for item in typescript["dimensionRecords"]],
            [(number, python_dimensions[number]["status"], python_dimensions[number]["verdict"]) for number in range(1, 38)],
        )

    def test_legacy_failed_review_remains_revise_in_python_and_typescript(self) -> None:
        review = {
            "score": 90,
            "passed": False,
            "issue_details": [{"dimension": 7, "severity": "warning", "description": "节奏偏慢"}],
        }
        python_manifest, _ = build_review_manifest(review, "ch_001", 70)
        script = """
import fs from 'node:fs'
import { buildDogReviewBundle } from './packages/openwrite-bridge/lib/dog-review.js'
const review = JSON.parse(fs.readFileSync(0, 'utf8'))
process.stdout.write(JSON.stringify(buildDogReviewBundle(review, 'ch_001', 70).manifest))
"""
        result = subprocess.run(
            ["node", "--input-type=module", "-e", script],
            cwd=REPO_ROOT,
            input=json.dumps(review),
            capture_output=True,
            text=True,
            check=True,
        )
        typescript_manifest = json.loads(result.stdout)

        self.assertEqual(python_manifest["gateStatus"], "pass")
        self.assertEqual(python_manifest["deliveryStatus"], "revise")
        self.assertEqual(typescript_manifest["deliveryStatus"], "revise")
        self.assertEqual(typescript_manifest["verdict"], "fail")

    def test_import_graph_checks_manifest_and_chapter(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            manuscript = root / "data" / "novels" / "book" / "data" / "manuscript" / "arc_001"
            manuscript.mkdir(parents=True)
            (manuscript / "ch_001.md").write_text("# 第一章\n\n正文", encoding="utf-8")
            source = root / "source.txt"
            source.write_text("第一章\n正文", encoding="utf-8")

            result = write_import_artifacts(
                root, "book", "IMPORT_TEST", source, "txt", "arc_001",
                [{"chapter_id": "ch_001", "title": "第一章", "writing_units": 2}],
                {"status": "completed", "summary": "无明显问题"},
            )
            graph_path = Path(result["graph_path"])
            artifact_dir = graph_path.parent
            self.assertEqual(_run_verifier("import-record", artifact_dir / "import.json")["verdict"], "pass")
            self.assertEqual(_run_verifier("import-chapter", manuscript / "ch_001.md")["verdict"], "pass")
            _validate_with_dog(graph_path, root)

    def test_delivery_requires_rereview_after_applied_revision(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            novel = root / "data" / "novels" / "book" / "data"
            manuscript = novel / "manuscript" / "arc_001" / "ch_001.md"
            manuscript.parent.mkdir(parents=True)
            manuscript.write_text("# 第一章\n\n初稿", encoding="utf-8")

            missing = write_delivery_artifacts(root, "book", "ch_001", 70)
            self.assertEqual(missing["stages"]["review"], "missing")
            self.assertFalse(missing["ready_for_delivery"])

            review_dir = novel / "reviews"
            review_dir.mkdir(parents=True)
            first_revision = "sha256:" + hashlib.sha256(manuscript.read_bytes()).hexdigest()
            review_path = review_dir / "ch_001.json"
            review_path.write_text(json.dumps({
                "score": 60,
                "passed": False,
                "source_revision": first_revision,
                "issue_details": [{"id": "issue-1", "severity": "blocker", "dimension": 1}],
            }), encoding="utf-8")
            failed = write_delivery_artifacts(root, "book", "ch_001", 70)
            self.assertEqual(failed["stages"]["closure"], "review_failed")

            review_path.write_text(json.dumps({
                "score": 90,
                "passed": False,
                "source_revision": first_revision,
                "issue_details": [{"id": "issue-warning", "severity": "warning", "dimension": 7}],
            }), encoding="utf-8")
            legacy_failed = write_delivery_artifacts(root, "book", "ch_001", 70)
            self.assertEqual(legacy_failed["stages"]["review"], "current")
            self.assertEqual(legacy_failed["stages"]["closure"], "review_failed")
            self.assertFalse(legacy_failed["ready_for_delivery"])

            review_path.write_text(json.dumps({
                "score": 60,
                "passed": False,
                "source_revision": first_revision,
                "issue_details": [{"id": "issue-1", "severity": "blocker", "dimension": 1}],
            }), encoding="utf-8")

            revision_dir = novel / "revisions" / "ch_001"
            revision_dir.mkdir(parents=True)
            proposal_path = revision_dir / "rev_test.json"
            proposal = {
                "proposal_id": "rev_test", "chapter_id": "ch_001", "kind": "review_fix",
                "status": "proposed", "source_revision": first_revision,
                "review_issue_ids": ["issue-1"],
            }
            proposal_path.write_text(json.dumps(proposal), encoding="utf-8")
            pending = write_delivery_artifacts(root, "book", "ch_001", 70)
            self.assertEqual(pending["stages"]["revision"], "proposal_pending")

            manuscript.write_text("# 第一章\n\n修订稿", encoding="utf-8")
            applied_revision = "sha256:" + hashlib.sha256(manuscript.read_bytes()).hexdigest()
            proposal.update({"status": "applied", "applied_revision": applied_revision})
            proposal_path.write_text(json.dumps(proposal), encoding="utf-8")
            stale_review = json.loads(review_path.read_text(encoding="utf-8"))
            stale_review.update({"stale": True, "stale_reason": "chapter_revised"})
            review_path.write_text(json.dumps(stale_review), encoding="utf-8")
            awaiting = write_delivery_artifacts(root, "book", "ch_001", 70)
            self.assertEqual(awaiting["stages"]["revision"], "applied_requires_rereview")
            self.assertEqual(awaiting["stages"]["application"], "applied")
            self.assertEqual(awaiting["stages"]["rereview"], "required")
            self.assertEqual(awaiting["stages"]["closure"], "rereview_required")
            self.assertFalse(awaiting["ready_for_delivery"])

            review_path.write_text(json.dumps({
                "score": 88,
                "passed": True,
                "source_revision": applied_revision,
                "issue_details": [],
                "issue_delta": {"resolved": [{"id": "issue-1"}], "remaining": [], "new": []},
            }), encoding="utf-8")
            closed = write_delivery_artifacts(root, "book", "ch_001", 70)
            self.assertEqual(closed["stages"]["closure"], "closed")
            self.assertTrue(closed["ready_for_delivery"])
            self.assertEqual(
                [edge["source"] for edge in closed["graph"]["dependsOn"]],
                ["review", "revision", "application", "rereview", "closure"],
            )
            self.assertFalse(any(
                node.get("verifier", {}).get("mode") == "agentic"
                for node in closed["graph"]["nodes"].values()
            ))
            _assert_acyclic(closed["graph"])
            delivery_dir = Path(closed["graph_path"]).parent
            self.assertEqual(_run_verifier("delivery-stage", delivery_dir / "closure.json")["verdict"], "pass")
            _validate_with_dog(Path(closed["graph_path"]), root)

    def test_docx_is_converted_before_text_decoding(self) -> None:
        try:
            from docx import Document
        except ImportError as error:  # pragma: no cover - dependency audit signal
            self.fail(f"python-docx is required by smart_import: {error}")
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "book.docx"
            document = Document()
            document.add_heading("第一章", level=1)
            document.add_paragraph("正文")
            document.save(path)
            text, detected = detect_and_convert(path)
            self.assertEqual(detected, "docx")
            self.assertIn("# 第一章", text)
            self.assertIn("正文", text)


if __name__ == "__main__":
    unittest.main()
